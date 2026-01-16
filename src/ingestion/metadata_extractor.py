"""
Document metadata extraction service.

Extracts basic metadata from documents before Gemini processing.
"""

import io
import logging
from pathlib import Path
from typing import Optional

from src.api.models.document import DocumentStructure, DocumentType

logger = logging.getLogger(__name__)


class DocumentMetadataExtractor:
    """
    Extracts metadata and text from documents.
    
    Supports PDF, DOCX, and TXT files.
    """

    # File extensions to MIME types
    MIME_TYPES = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
    }

    # Keywords for document type detection
    TYPE_KEYWORDS = {
        DocumentType.SOW: [
            "statement of work",
            "scope of work",
            "sow",
            "deliverables",
            "milestones",
            "project scope",
        ],
        DocumentType.CONTRACT: [
            "agreement",
            "contract",
            "terms and conditions",
            "binding agreement",
            "hereby agrees",
            "parties agree",
        ],
        DocumentType.MSA: [
            "master service agreement",
            "master services agreement",
            "msa",
            "framework agreement",
        ],
        DocumentType.AMENDMENT: [
            "amendment",
            "addendum",
            "modification",
            "change order",
        ],
        DocumentType.EMAIL: [
            "from:",
            "to:",
            "subject:",
            "sent:",
            "cc:",
            "re:",
        ],
    }

    # Keywords for structure detection
    STRUCTURE_KEYWORDS = {
        DocumentStructure.STRUCTURED: [
            "section",
            "article",
            "clause",
            "1.",
            "2.",
            "a)",
            "b)",
            "table of contents",
        ],
        DocumentStructure.SEMI_STRUCTURED: [
            "regarding",
            "as discussed",
            "please find",
            "attached",
        ],
    }

    def extract_text_from_pdf(self, content: bytes) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            content: PDF file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            raise

    def extract_text_from_docx(self, content: bytes) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            content: DOCX file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(content))
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise

    def extract_text(self, content: bytes, filename: str) -> str:
        """
        Extract text from a document based on file type.
        
        Args:
            content: File content as bytes
            filename: Original filename
            
        Returns:
            Extracted text
        """
        extension = Path(filename).suffix.lower()
        
        if extension == ".pdf":
            return self.extract_text_from_pdf(content)
        elif extension == ".docx":
            return self.extract_text_from_docx(content)
        elif extension == ".txt":
            return content.decode("utf-8", errors="replace")
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    def detect_document_type(self, text: str) -> DocumentType:
        """
        Detect document type based on content.
        
        Args:
            text: Document text content
            
        Returns:
            Detected document type
        """
        text_lower = text.lower()
        scores = {doc_type: 0 for doc_type in DocumentType}
        
        for doc_type, keywords in self.TYPE_KEYWORDS.items():
            for keyword in keywords:
                if keyword.lower() in text_lower:
                    scores[doc_type] += 1
        
        # Find the type with highest score
        max_score = max(scores.values())
        if max_score > 0:
            for doc_type, score in scores.items():
                if score == max_score:
                    return doc_type
        
        return DocumentType.OTHER

    def detect_structure(self, text: str) -> DocumentStructure:
        """
        Detect document structure based on content.
        
        Args:
            text: Document text content
            
        Returns:
            Detected structure type
        """
        text_lower = text.lower()
        
        # Check for structured indicators
        structured_count = 0
        for keyword in self.STRUCTURE_KEYWORDS[DocumentStructure.STRUCTURED]:
            if keyword.lower() in text_lower:
                structured_count += 1
        
        # Check for semi-structured indicators
        semi_count = 0
        for keyword in self.STRUCTURE_KEYWORDS[DocumentStructure.SEMI_STRUCTURED]:
            if keyword.lower() in text_lower:
                semi_count += 1
        
        # Determine structure based on scores
        if structured_count >= 3:
            return DocumentStructure.STRUCTURED
        elif semi_count >= 2 or structured_count >= 1:
            return DocumentStructure.SEMI_STRUCTURED
        else:
            return DocumentStructure.UNSTRUCTURED

    def extract_metadata(
        self,
        content: bytes,
        filename: str,
    ) -> dict:
        """
        Extract all metadata from a document.
        
        Args:
            content: File content as bytes
            filename: Original filename
            
        Returns:
            Dictionary with extracted metadata
        """
        # Extract text
        text = self.extract_text(content, filename)
        
        # Detect type and structure
        document_type = self.detect_document_type(text)
        structure_type = self.detect_structure(text)
        
        return {
            "raw_text": text,
            "raw_text_length": len(text),
            "document_type": document_type,
            "structure_type": structure_type,
            "filename": filename,
            "file_size": len(content),
            "content_type": self.MIME_TYPES.get(
                Path(filename).suffix.lower(), "application/octet-stream"
            ),
        }


def get_metadata_extractor() -> DocumentMetadataExtractor:
    """Get a metadata extractor instance."""
    return DocumentMetadataExtractor()
